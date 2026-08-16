from pathlib import Path
import sys, json, csv, re, math
sys.path.insert(0, str(Path(__file__).parent))
import build_mas_reliab as assets

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path('/home/user/mas_reliab'); FIG=ROOT/'figures'; SUP=ROOT/'support'
NAVY='0B1F33'; BLUE='146C94'; TEAL='1B998B'; AMBER='F4A261'; RED='D95D5D'; LIGHT='F3F6F8'; MID='DCE5EA'; DARK='23323D'; GRAY='657681'; PURPLE='7A5AF8'


def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def set_cell_text_color(cell, color):
    for p in cell.paragraphs:
        for r in p.runs: r.font.color.rgb=RGBColor.from_string(color)

def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); e=OxmlElement('w:tblHeader'); e.set(qn('w:val'),'true'); trPr.append(e)

def cant_split(row):
    trPr=row._tr.get_or_add_trPr(); e=OxmlElement('w:cantSplit'); trPr.append(e)

def add_field(paragraph, instruction):
    run=paragraph.add_run(); fldChar=OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=instruction
    fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'separate')
    fldChar3=OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'),'end')
    run._r.extend([fldChar,instr,fldChar2,fldChar3])

def setup_doc():
    d=Document(); sec=d.sections[0]
    sec.top_margin=Inches(.75); sec.bottom_margin=Inches(.72); sec.left_margin=Inches(.82); sec.right_margin=Inches(.82)
    sec.page_width=Inches(8.27); sec.page_height=Inches(11.69)
    styles=d.styles
    normal=styles['Normal']; normal.font.name='Times New Roman'; normal.font.size=Pt(10.3)
    normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.08
    for i,size,color in [(1,16,NAVY),(2,13,BLUE),(3,11,TEAL),(4,10,DARK)]:
        s=styles[f'Heading {i}']; s.font.name='Aptos Display'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(11); s.paragraph_format.space_after=Pt(5); s.paragraph_format.keep_with_next=True
    styles['Title'].font.name='Aptos Display'; styles['Title'].font.color.rgb=RGBColor.from_string(NAVY)
    styles['Caption'].font.name='Aptos'; styles['Caption'].font.size=Pt(9); styles['Caption'].font.italic=True; styles['Caption'].font.color.rgb=RGBColor.from_string(GRAY)
    if 'Equation' not in [s.name for s in styles]:
        es=styles.add_style('Equation',WD_STYLE_TYPE.PARAGRAPH); es.font.name='Cambria Math'; es.font.size=Pt(10.5); es.paragraph_format.space_before=Pt(4); es.paragraph_format.space_after=Pt(6); es.paragraph_format.keep_together=True
    if 'Callout' not in [s.name for s in styles]:
        cs=styles.add_style('Callout',WD_STYLE_TYPE.PARAGRAPH); cs.font.name='Aptos'; cs.font.size=Pt(9.5); cs.font.color.rgb=RGBColor.from_string(DARK); cs.paragraph_format.left_indent=Inches(.22); cs.paragraph_format.right_indent=Inches(.22); cs.paragraph_format.space_before=Pt(6); cs.paragraph_format.space_after=Pt(7)
    # Header/footer
    h=sec.header.paragraphs[0]; h.text='MAS-RELIAB  |  Research Paper and Viva  |  Dev Parth'; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in h.runs: r.font.name='Aptos'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
    f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.add_run('Page '); add_field(f,'PAGE'); f.add_run(' of '); add_field(f,'NUMPAGES')
    for r in f.runs: r.font.name='Aptos'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
    return d

def para(d,text='',bold_lead=None,italic=False,style=None,align=None):
    p=d.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold=True; p.add_run(text[len(bold_lead):])
    else: p.add_run(text)
    if italic:
        for r in p.runs: r.italic=True
    if align is not None: p.alignment=align
    return p

def bullets(d,items,level=0):
    for item in items:
        p=d.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.paragraph_format.space_after=Pt(2)
        if isinstance(item,tuple):
            r=p.add_run(item[0]); r.bold=True; p.add_run(item[1])
        else: p.add_run(item)

def numbered(d,items):
    for item in items:
        p=d.add_paragraph(style='List Number'); p.paragraph_format.space_after=Pt(3); p.add_run(item)

def equation(d,text,label=None):
    p=d.add_paragraph(style='Equation'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(text)
    if label: p.add_run('    '+label)
    return p

def table(d,headers,rows,widths=None,small=8.6):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=t.rows[0]; repeat_table_header(hdr)
    for i,h in enumerate(headers):
        c=hdr.cells[i]; c.text=str(h); shade(c,NAVY); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c)
        for p in c.paragraphs: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in c.paragraphs[0].runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.name='Aptos'; r.font.size=Pt(small)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells; cant_split(t.rows[-1])
        for i,val in enumerate(row):
            cells[i].text=str(val); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP; set_cell_margins(cells[i])
            if ri%2==1: shade(cells[i],'F7F9FA')
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(1)
                for r in p.runs: r.font.name='Aptos'; r.font.size=Pt(small)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    d.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def caption(d,text):
    p=d.add_paragraph(text,style='Caption'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; return p

def figure(d,file,cap,width=6.55):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(FIG/file),width=Inches(width)); caption(d,cap)

def page_break(d): d.add_page_break()

def callout(d,title,text,color=TEAL):
    t=d.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,'EAF5F3' if color==TEAL else 'FFF3E7'); set_cell_margins(c,110,150,110,150)
    p=c.paragraphs[0]; r=p.add_run(title+'  '); r.bold=True; r.font.color.rgb=RGBColor.from_string(color); p.add_run(text)
    for r in p.runs: r.font.name='Aptos'; r.font.size=Pt(9.5)
    d.add_paragraph().paragraph_format.space_after=Pt(1)

d=setup_doc()
# Title page
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(70)
r=p.add_run('A FRAMEWORK FOR EVALUATING\nRELIABILITY AND FAILURE PROPAGATION\nIN MULTI-AGENT AI SYSTEMS'); r.font.name='Aptos Display'; r.font.size=Pt(23); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('MAS-RELIAB'); r.font.name='Aptos Display'; r.font.size=Pt(19); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(TEAL)
para(d,'Multi-Agent System Reliability and Failure Propagation Evaluation Framework',italic=True,align=WD_ALIGN_PARAGRAPH.CENTER)
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(36); r=p.add_run('Dev Parth'); r.bold=True; r.font.name='Aptos'; r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(BLUE)
para(d,'Research Paper and Viva Defence Package',align=WD_ALIGN_PARAGRAPH.CENTER)
para(d,'Pre-experimental, results-ready manuscript  •  16 August 2026',align=WD_ALIGN_PARAGRAPH.CENTER)
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(42)
r=p.add_run('SCIENTIFIC STATUS'); r.bold=True; r.font.name='Aptos'; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(RED)
para(d,'No experiment logs or empirical data were supplied. Consequently, all Results fields are intentionally marked with em dashes and every empirical claim remains pending. This document specifies a complete protocol and reporting shell; it does not fabricate findings.',italic=True,align=WD_ALIGN_PARAGRAPH.CENTER)
page_break(d)

# Front matter
p=d.add_heading('Abstract',level=1)
para(d,'Multi-agent AI systems can distribute reasoning, tool use, and decision-making across specialized agents, but they can also convert a local error into a system-level failure through message and state dependencies. Conventional task-success scores provide weak diagnostic evidence: they do not show whether behaviour is repeatable, where a fault began, whether it propagated, whether an observer can attribute it, how a mitigation changed the cascade, or what reliability cost was incurred. This paper specifies MAS-RELIAB, a unified experimental methodology for evaluating these dimensions under controlled and reproducible conditions. An episode is represented as a directed dependency graph augmented with time-ordered trace events, a documented fault intervention, deterministic or state-based outcome checks, and resource records. MAS-RELIAB reports a non-aggregated reliability vector spanning task success, repeated-run consistency, perturbation and fault robustness, propagation containment, attribution accuracy, and resource use. The protocol compares single-agent, sequential, parallel, and hierarchical topologies; manipulates fault type and topological position; tests output-only, partial, and full observability; and evaluates verification and recovery strategies. AI Command Center is treated as the execution and observability substrate, while MAS-RELIAB supplies benchmark tasks, injection operators, metrics, experimental controls, and analysis. Six preregistered experiments and an ablation plan map directly to RQ1–RQ6 and H1–H6. Because no executions are available at authoring time, this manuscript is deliberately pre-experimental: Results tables contain no values, and the conclusion claims only that the protocol has been specified. The contribution is therefore not a priority claim for graph-based cascade analysis, but a results-ready methodology connecting reliability, fault injection, propagation, observability-conditioned attribution, mitigation, topology, and cost within one auditable design.')
para(d,'Keywords: multi-agent AI; reliability evaluation; failure propagation; fault injection; failure attribution; observability; agent topology; mitigation; reproducibility')

callout(d,'Integrity note.','The Results section may be populated only from signed run manifests and analysis outputs. Hypotheses are not promises: contradictory, null, or mixed findings must be reported as observed.',RED)

d.add_heading('Document map',level=1)
table(d,['Part','Purpose'],[
 ('Research manuscript','Sections 1–15: motivation, related work, formalism, MAS-RELIAB, benchmark design, six experiments, analysis, results shell, discussion, limitations, and conclusion.'),
 ('Appendices','Operational definitions, schemas, pseudocode, reproducibility checklist, statistical decision rules, and empty reporting tables.'),
 ('Viva defence pack','Opening statement, timed presentation script, slide map, examiner questions with model answers, and defence boundaries.'),
 ('Companion deck','MAS_RELIAB_Viva_Deck_Dev_Parth.pptx; the deck mirrors this manuscript and labels empirical slides as pending.')],widths=[1.5,5.0])

page_break(d)
d.add_heading('Contents',level=1)
para(d,'Update this field in Microsoft Word (right-click → Update Field) if section pagination changes.',italic=True)
p=d.add_paragraph(); add_field(p,'TOC \\o "1-3" \\h \\z \\u')
page_break(d)

# 1 Introduction
d.add_heading('1. Introduction',level=1)
para(d,'Large language models are increasingly organized as teams rather than isolated responders. A planner may decompose a task, specialists may call tools, a reviewer may check intermediate work, and a manager may assemble the final state. This organization can improve modularity and coverage, but it also introduces dependency edges: an incorrect premise, malformed tool output, stale observation, or ambiguous handoff may be accepted by downstream agents and amplified into a coherent yet wrong answer. The operational question is therefore not merely “Did the system succeed?” but “How reliably did it succeed, what failed when it did not, how far did the failure travel, could it be attributed, could it be contained, and at what cost?”')
para(d,'Single-run task success is an incomplete answer. A system may pass once but fail under another seed; achieve the correct text while leaving the environment in the wrong state; appear robust because an injected fault was never consumed; or recover through repeated expensive calls. Conversely, a failed final answer may result from one decisive early error, several independent errors, or a correct plan defeated by an environment fault. Treating these cases as the same binary outcome obscures engineering decisions and weakens scientific comparison.')
para(d,'Recent research covers important portions of this space. MAST provides an empirical taxonomy and annotated traces for multi-agent failure analysis [1]. MultiAgentBench studies collaboration and multiple coordination topologies [2]. τ-bench motivates state-based evaluation and repeated-run pass^k [3]. ReliabilityBench introduces production-like perturbation and tool/API fault stress for agents [4]. TraceElephant and AgenTracer focus on agent- and step-level failure attribution under trace evidence and programmed faults [5,6], while ErrorProbe develops backward, evidence-grounded diagnosis [7]. AgentAsk analyzes errors at handoffs and uses clarification to contain them [8]. RiskLab, TAMAS, and SILO-BENCH address controlled emergent risks, adversarial safety, and distributed coordination [9–11]. Most importantly for positioning, From Spark to Fire already models multi-agent collaboration as a directed dependency graph, injects atomic error seeds, studies cascade and topology effects, and evaluates a governance layer [12]. MAS-RELIAB therefore does not claim that graph-based propagation or cascade mitigation is new.')
para(d,'The defensible contribution is a unified, auditable experimental protocol that connects these concerns in the same paired episodes: multidimensional reliability; controlled and verified fault injection; graph-normalized propagation; observability-conditioned attribution; verification and recovery comparisons; budget-matched topology controls; and resource trade-offs. This scope is methodological rather than a claim that every component is individually unprecedented.')

d.add_heading('1.1 Aim',level=2)
para(d,'The aim is to design and empirically validate MAS-RELIAB as a reproducible method for measuring reliability and failure propagation in multi-agent AI systems, while keeping execution infrastructure separate from benchmark and analysis logic.')

d.add_heading('1.2 Objectives',level=2)
numbered(d,[
 'Define a directed, time-aware representation of agents, dependencies, trace events, faults, outcomes, and interventions.',
 'Operationalize a failure taxonomy and deterministic fault operators whose application and consumption can be verified.',
 'Report reliability as interpretable components before any optional composite or ranking.',
 'Measure propagation, attribution, containment, recovery, and cost under paired controls.',
 'Compare single-agent, sequential, parallel, and hierarchical architectures without silently confounding topology with model, task, or compute budget.',
 'Provide reproducible schemas, preregistered hypotheses, statistical decision rules, and results-ready artifacts.'
])

d.add_heading('1.3 Contributions',level=2)
bullets(d,[
 ('Unified protocol, ','A common episode design links task outcomes to verified faults, dependency graphs, trace evidence, mitigation, and resource accounting.'),
 ('Operational measurement, ','Definitions distinguish event propagation, normalized downstream impact, amplification, depth, edge transmission, detection, attribution, and recovery.'),
 ('Controlled comparisons, ','The same task/seed blocks support causal comparisons across positions, observability views, verification regimes, recovery strategies, and topologies.'),
 ('Infrastructure separation, ','AI Command Center executes and observes; MAS-RELIAB defines benchmark tasks, interventions, evaluators, metrics, and analysis.'),
 ('Results integrity, ','Empty tables, run manifests, analysis locks, and explicit decision rules prevent invented or selectively reported outcomes.'),
 ('Viva readiness, ','The document includes a defence script and answers to likely questions about novelty, causality, statistics, validity, and limitations.')
])

d.add_heading('1.4 Scope and non-claims',level=2)
callout(d,'Non-claim.','MAS-RELIAB is not presented as the first framework for multi-agent reliability, graph-based cascade analysis, fault injection, failure attribution, or mitigation. Its intended value lies in the integration and controlled experimental linkage of these dimensions.',RED)
para(d,'The initial scope is tool-using language-agent systems whose interactions and environment states can be logged. The protocol does not establish reliability for embodied robots, open-world human teams, continuously learning agents, or safety-critical deployment without additional domain validation. The present manuscript also does not establish empirical superiority because no MAS-RELIAB runs have yet been supplied.')

# 2 Related work
d.add_heading('2. Related Work and Positioning',level=1)
d.add_heading('2.1 Multi-agent failure taxonomies and coordination benchmarks',level=2)
para(d,'MAST derives fourteen failure modes from annotated traces across several multi-agent frameworks and organizes them into system-design, inter-agent-alignment, and verification groups [1]. MAS-RELIAB uses MAST as a conceptual anchor but adds intervention metadata, propagation measurements, and experimental comparisons; it does not replace MAST. MultiAgentBench evaluates task completion and collaboration under star, chain, tree, and graph protocols [2], making topology a legitimate independent variable. SILO-BENCH further shows why distributed information exchange and successful information integration should be distinguished [11].')

d.add_heading('2.2 Reliability, state-based evaluation, and controlled stress',level=2)
para(d,'τ-bench evaluates tool-agent-user tasks by comparing the final database state to an annotated goal state and introduces pass^k for repeated trials [3]. MAS-RELIAB adopts the principles of end-state evaluation and repeated execution. ReliabilityBench evaluates consistency, paraphrase robustness, and controlled tool/API faults under production-like stress [4]. MAS-RELIAB extends this logic to multi-agent dependency structures, attribution, topology, propagation, and mitigation, while retaining state equivalence wherever possible. Surveys of agent evaluation emphasize reproducibility, environment control, cost, robustness, and behaviour-level evidence [13,14]. A later unified agent-evaluation framework standardizes instruction–tool–environment execution and separates benchmark outcomes from scaffold and environment effects [15], reinforcing the need to pin infrastructure and environment versions.')

d.add_heading('2.3 Attribution, observability, and intervention',level=2)
para(d,'TraceElephant studies attribution with full execution traces and reproducible environments [5]. AgenTracer uses counterfactual replay and programmed fault injection to curate agent- and step-level origin labels [6]. ErrorProbe operationalizes taxonomy-guided anomaly detection and backward tracing with executable validation [7]. These works make two points central to MAS-RELIAB: attribution accuracy must be evaluated against explicit origin labels, and the available evidence view is itself an experimental condition. AgentAsk identifies handoff-level error categories and evaluates targeted clarification as a containment mechanism [8]. RiskLab provides controlled topology–environment–protocol–agent–task configurations for emergent-risk studies [9], while TAMAS supplies multi-agent adversarial scenarios and safety/effectiveness trade-offs [10].')

d.add_heading('2.4 Error cascades and novelty boundary',level=2)
para(d,'From Spark to Fire is the closest overlap: it abstracts collaboration as a directed dependency graph, studies cascade amplification, topological sensitivity and consensus inertia, injects atomic errors, and evaluates a message-layer governance mechanism [12]. MAS-RELIAB must therefore be evaluated on a narrower claim: whether one protocol can connect multidimensional reliability with controlled interventions, graph-normalized propagation, observability-conditioned attribution, multiple verification/recovery baselines, budget-matched topologies, and cost. A successful MAS-RELIAB paper should demonstrate the utility of that linkage rather than imply priority for any isolated ingredient.')

caption(d,'Table 1. Scope comparison based on the primary dimensions explicitly reported in the cited work. “Partial” means the dimension is present but not the central evaluation target; “not central” is not a claim that the implementation lacks the capability. Recheck every cell against the final paper version before submission.')
rows=[
 ('MAST [1]','Reported','Not central','Not central','Failure dynamics / taxonomy','Trace annotation','Not central','Not central'),
 ('MultiAgentBench [2]','Reported','Not central','Reported','Coordination KPIs','Not central','Strategy comparisons','Reported'),
 ('τ-bench [3]','State-based + pass^k','Not central','Single-agent focus','Not central','Not central','Not central','Repeated trials'),
 ('ReliabilityBench [4]','State equivalence + repeats','Tool/API faults','Not central','Fault tolerance','Not central','Architecture stress','Cost reported'),
 ('TraceElephant [5]','Failure traces','Not central','Configurations','Not central','Full vs partial trace','Attribution methods','Not central'),
 ('AgenTracer [6]','Failed trajectories','Programmed faults','Agentic systems','Partial','Agent + step','Feedback / correction','Not central'),
 ('AgentAsk [8]','Five benchmarks','Handoff corruption','Architecture-agnostic','Edge propagation','Critical handoff','Clarification','Latency + cost'),
 ('RiskLab [9]','Trajectory-grounded','Controlled risks','Flexible topology','Emergent risk','Detector registry','Protocol comparison','Configurable'),
 ('TAMAS [10]','Safety + effectiveness','Adversarial attacks','Multiple configurations','Attack effects','Not central','Defence need','Trade-off score'),
 ('SILO-BENCH [11]','Algorithmic tasks','Not central','Scale / communication','Integration failure','Stage localization','Not central','Tokens / density'),
 ('From Spark to Fire [12]','Framework tasks','Atomic error seed','Topological sensitivity','Core focus','Genealogy evidence','Governance plugin','Not central'),
 ('MAS-RELIAB','State-based + repeats','Verified operators','Four baselines','Core metrics','Evidence-view experiment','Verification + recovery','Pareto analysis')]
table(d,['Work','Outcome / repeat','Controlled faults','Topology','Propagation','Attribution / evidence','Mitigation','Resources'],rows,widths=[1.25,.82,.78,.70,.80,.88,.82,.72],small=7.2)
figure(d,'figure_01_architecture.png','Figure 1. Separation of concerns: MAS-RELIAB is the methodology and analysis layer; AI Command Center is the execution and observability substrate.')

# 3 Questions/hypotheses
d.add_heading('3. Research Questions and Hypotheses',level=1)
para(d,'The hypotheses are preregistered directional or interaction claims. They must be tested against the specified primary metrics; unsupported hypotheses will be rejected or left inconclusive rather than reworded after inspection.')
rq_rows=[
 ('RQ1','How should reliability be measured beyond a single task-success score?','H6','The multidimensional profile will reveal rank reversals or masked weaknesses not visible in task success alone.','E1'),
 ('RQ2','How do fault type, severity, and injection position affect propagation?','H1','Earlier topological injections will produce greater reachable-set-normalized impact and depth than late injections, conditional on verified consumption.','E2'),
 ('RQ3','How does execution observability affect failure attribution?','H2','Full traces will improve agent- and step-level attribution over partial and output-only views on the same failed episodes.','E3'),
 ('RQ4','Does local verification contain cascades better than final-only checking?','H3','Local verification will reduce event propagation and downstream affected fraction relative to final-only and no verification.','E4'),
 ('RQ5','Which recovery strategy offers the best reliability–cost trade-off?','H4','Rollback and independent alternate-agent recovery will recover more faults than same-context retry, but with additional resource cost.','E5'),
 ('RQ6','How does topology alter reliability and propagation?','H5','Topology will interact with fault location: parallel aggregation will contain isolated worker faults more effectively, while hierarchical manager faults will have larger normalized impact than leaf faults.','E6')]
table(d,['Research question','Question','Hyp.','Preregistered hypothesis','Experiment'],rq_rows,widths=[.58,2.55,.45,2.80,.55],small=7.9)
para(d,'H6 is tested through prespecified discordance analyses: architecture rankings by task success will be compared with rankings by pass^k, fault degradation, propagation containment, attribution accuracy, latency, tokens, and cost. A “masked weakness” requires a statistically and practically meaningful disadvantage on at least one preregistered component, not a post-hoc narrative.')

# 4 Formal model
d.add_heading('4. Problem Formulation',level=1)
d.add_heading('4.1 Episode and graph',level=2)
para(d,'A MAS-RELIAB episode is the tuple 𝓔 = ⟨τ, G, Π, Ω, I, X, Y, C⟩. Here τ is a versioned task; G=(V,E) is a directed dependency graph; Π contains model, role, prompt, tool, routing, and budget policies; Ω is the environment snapshot; I is either a null intervention or a documented fault intervention; X is the time-ordered execution trace; Y is the evaluated outcome; and C is the resource record.')
equation(d,'G = (V, E),   V = {agents, tools, state objects, decision events},   E ⊆ V × V × T','(1)')
para(d,'An edge (u,v,t) means that event or node v consumed information, state, or control from u at time t. Agent-only projections are used for architecture comparison; event-level graphs are used for step attribution. Edges are created from message IDs, tool-call parent IDs, state reads/writes, explicit routes, and aggregator provenance. Inferred semantic dependencies are stored separately with a confidence label and are not mixed with instrumented edges in the primary analysis.')

figure(d,'figure_02_graph.png','Figure 2. Example dependency graph with a verified injected origin, downstream failures, and a contained branch. This is a conceptual diagram, not an empirical result.')

d.add_heading('4.2 Fault intervention',level=2)
para(d,'A fault intervention I = ⟨f, v*, t*, s, z, pre, post⟩ identifies a fault type f, target v*, trigger time or event t*, severity s, random seed z, a precondition, and postconditions. Two checks are required: applied(I), showing that the operator changed the intended artifact, and consumed(I), showing that the targeted agent or downstream component received the changed artifact. Episodes failing either check are injection failures and are reported separately; they do not enter conditional propagation estimates.')
equation(d,'I_valid = 𝟙[applied(I)=1 ∧ consumed(I)=1]','(2)')
para(d,'Fault payloads must be task-bounded and reproducible. For example, a tool schema-drift operator changes a pinned field name according to a versioned patch, while a reasoning-seed operator replaces one validated intermediate fact with a predeclared alternative. Free-form instructions such as “make the agent fail” are prohibited because they do not define dose, target, or repeatability.')

d.add_heading('4.3 Outcome and ground truth',level=2)
para(d,'Primary task correctness is determined by a deterministic end-state validator whenever feasible: database equality under permitted equivalence, unit tests, constraint satisfaction, exact structured answers, or fact keys against a closed evidence pack. Text similarity and an LLM judge may be secondary diagnostics, never the sole ground truth for the core benchmark. An episode may satisfy some constraints and not others; therefore validators return both a binary all-required-constraints result and a vector of constraint-level outcomes.')
equation(d,'Success(𝓔) = 𝟙[∀c ∈ C_required : c(Y, Ω_final)=1]','(3)')
para(d,'Fault-origin ground truth combines the intervention log with trace evidence. For naturally occurring failures, two trained annotators independently label the earliest decisive error event and responsible agent using a published codebook. Disagreements are adjudicated by a third reviewer. The primary attribution experiment uses injected or replay-validated cases because their origins are more objectively checkable.')

d.add_heading('4.4 Reliability vector',level=2)
para(d,'MAS-RELIAB treats reliability as an unaggregated profile rather than a single number. For condition c, the reported vector is:')
equation(d,'R(c) = [TSR ↑, pass^k ↑, Consistency ↑, Δpert ↓, Δfault ↓, EPR ↓, DAF ↓, Attribution ↑, Recovery ↑, Tokens ↓, Latency ↓, Cost ↓]','(4)')
para(d,'The arrows indicate preferred direction, not commensurate units. Individual components, denominators, confidence intervals, and missingness are reported first. Any optional scalar is explicitly labelled exploratory, accompanied by the weighting rule and a sensitivity analysis across plausible weights. No arbitrary weighting is presented as scientifically justified.')

figure(d,'figure_06_metrics.png','Figure 3. MAS-RELIAB reports interpretable metric families separately and uses Pareto analysis for trade-offs.')

d.add_heading('4.5 Propagation model',level=2)
para(d,'Let D_i be the set of nodes reachable downstream from the injected origin i in the instrumented graph, and A_i ⊆ D_i the nodes that exhibit a new task-relevant erroneous state after consuming the fault lineage. A node is not marked affected merely because it appears later in the trace; a validator or coded causal rule must link its state to the injected lineage.')
equation(d,'EPR = (Σ_e 𝟙[|A_i(e)| > 0]) / N_valid_injections','(5)')
equation(d,'DAF(e) = |A_i(e)| / |D_i(e)|,  with DAF(e)=0 when |D_i(e)|=0','(6)')
equation(d,'AF(e) = |A_i(e)| / max(1, |O_i(e)|),   Depth(e)=max_{v∈A_i} dist_G(i,v)','(7)')
para(d,'EPR is event-level propagation: the proportion of valid injections that cause any downstream failure. DAF normalizes for how many downstream nodes were actually reachable, reducing the trivial advantage of late injections. AF is an amplification factor relative to the set O_i of initially corrupted nodes. Depth is the maximum directed shortest-path distance. Edge transmission p_uv is estimated only on opportunities where u was erroneous and v consumed u before independently failing.')
equation(d,'p̂_uv = N(u erroneous → v newly erroneous after consumption) / N(v consumed erroneous u)','(8)')
para(d,'If edges transmit independently, the probability of traversal along one serial path π is ∏_{e∈π}p_e. For multiple edge-disjoint paths, 1−∏_π(1−∏_{e∈π}p_e) is an approximation. Because multi-agent messages are rarely independent, MAS-RELIAB does not use that approximation as primary evidence; empirical cascade probabilities are estimated by paired interventions or Monte Carlo replay, with assumptions stated.')

d.add_heading('4.6 Attribution targets',level=2)
para(d,'The attribution target has two levels: responsible agent a* and earliest decisive step t*. “Decisive” means that a counterfactual repair at that event, with the preceding prefix held fixed and downstream execution replayed under a prespecified policy, restores a feasible path to success or removes the labelled failure lineage. Metrics include exact agent accuracy, exact step accuracy, joint accuracy, top-k recall, mean reciprocal rank, and absolute/topological localization distance. Multiple-root cases receive a set-valued gold label and are scored with set precision and recall.')

# 5 Framework
d.add_heading('5. MAS-RELIAB Framework',level=1)
d.add_heading('5.1 Separation from AI Command Center',level=2)
para(d,'AI Command Center is the execution and observability infrastructure. It launches episodes, routes agent and tool calls, captures prompts and responses according to governance policy, records state changes and timestamps, and stores run artifacts. MAS-RELIAB is the benchmark and analysis layer: it defines tasks, architecture adapters, intervention operators, gold validators, graph construction, metrics, statistical comparisons, and reporting. This boundary allows the protocol to be implemented on another orchestration platform without changing its scientific definitions.')
infra_rows=[
 ('Episode execution','AI Command Center','Execute a frozen configuration and return an immutable run ID.'),
 ('Trace and artifact capture','AI Command Center','Store model/tool I/O, messages, state deltas, timing, errors, and hashes.'),
 ('Task and gold-state catalog','MAS-RELIAB','Version tasks, environments, constraints, and allowed equivalences.'),
 ('Fault injection','MAS-RELIAB through adapter','Apply typed operators; verify application and consumption.'),
 ('Graph and lineage construction','MAS-RELIAB','Build instrumented dependencies; preserve inferred edges separately.'),
 ('Evaluation and statistics','MAS-RELIAB','Compute metrics, paired analyses, uncertainty, and reporting artifacts.')]
table(d,['Responsibility','Owner','Required contract'],infra_rows,widths=[1.45,1.35,3.75],small=8.3)

figure(d,'figure_01_architecture.png','Figure 4. Framework architecture. The repeated figure is included here to anchor the implementation boundary.')

d.add_heading('5.2 Required adapter interface',level=2)
para(d,'A conforming execution substrate exposes four logical interfaces, regardless of programming language:')
bullets(d,[
 ('execute_episode(config) → run_id, ','launch exactly one pinned episode and return a durable identifier.'),
 ('stream_or_export_trace(run_id) → events, ','return ordered, schema-valid events with parent and provenance identifiers.'),
 ('snapshot_environment(run_id, point) → digest, ','record pre-run, checkpoint, and final environment hashes or snapshots.'),
 ('retrieve_artifacts(run_id) → manifest, ','return prompts, versions, injection record, validator output, resource use, and integrity hashes.')
])
para(d,'The framework rejects episodes missing required provenance fields. Sensitive internal reasoning that cannot be lawfully logged is not assumed to be available; such systems can participate in output-only or partial-trace conditions, but their observability level must be declared.')

d.add_heading('5.3 Event trace schema',level=2)
para(d,'Each event has a run ID, event ID, parent event IDs, actor ID, role, event type, timestamp, model/tool version, input artifact hashes, output artifact hash, state delta, token and latency fields, visibility policy, fault-lineage tags, and validation annotations. The full JSON schema appears in Appendix B and as a companion file.')
figure(d,'figure_07_observability.png','Figure 5. Evidence views for the paired attribution experiment. No accuracy values are shown because results are pending.')

d.add_heading('5.4 Fault taxonomy and operators',level=2)
para(d,'The operational taxonomy has five families: agent-local, handoff/edge, tool/environment, coordination, and verification/recovery. It is deliberately cross-referenced to MAST [1] and AgentAsk [8]. Categories help stratify analysis; the operator definition, not the category name, supplies reproducibility.')
operator_rows=[
 ('F1: Incorrect intermediate fact','Agent-local','Replace one validated intermediate value with a predeclared wrong value of matched format.','Value differs; checksum logged.','Target reads the altered value.'),
 ('F2: Role non-compliance','Agent-local','Swap one role constraint for a versioned contradictory instruction.','Prompt diff matches patch.','Agent invocation uses patched prompt.'),
 ('F3: Message truncation','Handoff','Remove a deterministic suffix or selected field at severity s.','Message hash and byte delta.','Receiver opens the mutated message.'),
 ('F4: Referential drift','Handoff','Change one entity/reference ID while preserving syntax.','Reference patch validated.','Receiver resolves the altered reference.'),
 ('F5: Tool corruption','Tool/environment','Return wrong, partial, stale, or schema-shifted response from a stub/snapshot.','Stub reports injected response ID.','Tool result enters agent context.'),
 ('F6: Timeout/rate limit','Tool/environment','Emit a deterministic error at selected call number.','Error code and trigger match.','Agent handles or propagates error.'),
 ('F7: Conflicting handoff','Coordination','Provide two validly formatted but mutually inconsistent upstream results.','Conflict constraints verified.','Coordinator consumes both.'),
 ('F8: Verification miss','Verification','Suppress or weaken one local check according to an ablation patch.','Validator hook disabled as declared.','Unchecked artifact is handed off.')]
table(d,['Operator','Family','Intervention','Applied check','Consumed check'],operator_rows,widths=[1.2,.9,2.35,1.1,1.1],small=7.3)
figure(d,'figure_03_taxonomy.png','Figure 6. Operational failure taxonomy. It organizes operators but does not supersede prior failure taxonomies.')

d.add_heading('5.5 Severity and position',level=2)
para(d,'Severity is defined by operator-specific, testable dose levels rather than adjectives. For truncation it may be a prespecified information fraction; for tool corruption it may be one field, several linked fields, or a complete response; for timeout it may be a single call or a bounded burst. Exact values are set during the pilot and then frozen. Position is measured two ways: normalized event time q=t/T and normalized topological depth d(i)/d_max. Primary strata are early, middle, and late terciles of topological depth. Analyses additionally include the size of the reachable downstream set to avoid mistaking opportunity for vulnerability.')

d.add_heading('5.6 Failure labelling workflow',level=2)
numbered(d,[
 'Validate task and environment before execution; reject a task if the clean oracle cannot satisfy the gold state.',
 'Apply the selected operator and store its patch, target, trigger, severity, seed, and before/after hashes.',
 'Verify application and consumption independently.',
 'Run deterministic local and end-state validators; construct the instrumented dependency graph.',
 'Assign propagated effects only when temporal order, dependency evidence, and a task-relevant error validator agree.',
 'For naturally occurring failures, perform blinded double annotation and adjudication.',
 'Freeze the labelled trace before attribution methods are evaluated.'
])

# 6 benchmark
d.add_heading('6. Benchmark Tasks and Ground Truth',level=1)
d.add_heading('6.1 Design principles',level=2)
bullets(d,[
 ('Objectively checkable outcomes. ','Core success is determined by environment state, executable tests, exact structured fields, or closed-corpus fact keys.'),
 ('Interaction necessity. ','Multi-agent variants must exchange information or state that is genuinely relevant; decorative role-play is excluded.'),
 ('Faultability. ','Each task exposes documented injection points with preconditions and independent checks.'),
 ('Topology portability. ','Equivalent task information is available to all four architectures under comparable budgets.'),
 ('Contamination control. ','Generated task instances, private splits, and versioned synthetic environments reduce memorization risk.'),
 ('Difficulty calibration. ','The clean baseline should avoid both floor and ceiling effects; unsolvable and trivial instances are revised or stratified.')
])

d.add_heading('6.2 Task families',level=2)
task_rows=[
 ('Closed-corpus research synthesis','Retrieve and reconcile evidence from a pinned document pack.','Claim table with required citations and contradiction flags.','Fact-key and evidence-span validator.'),
 ('Data analysis','Inspect versioned CSV/JSON data and compute requested statistics.','Structured results plus generated artifact.','Reference computation with tolerances and schema checks.'),
 ('Software debugging','Diagnose and patch a sandboxed repository snapshot.','Patch and test log.','Hidden and public unit tests; lint/security constraints.'),
 ('Constraint-based decision support','Select or rank options under explicit hard and soft constraints.','Decision object with constraint trace.','Constraint solver and exact feasibility checks.'),
 ('Stateful tool workflow','Use APIs to update a simulated service.','Final database and action log.','End-state equivalence and policy-rule validator.')]
table(d,['Task family','Core activity','Output','Primary ground truth'],task_rows,widths=[1.35,1.75,1.55,1.85],small=8.0)

d.add_heading('6.3 Instance schema',level=2)
para(d,'Every task instance records task_id, family, version, natural-language instruction, initial environment snapshot, permitted tools, role information policy, gold constraints, equivalence rules, timeout, budget, injection points, clean-oracle evidence, sensitivity class, and license/provenance. The benchmark split is fixed before the final run. A task is excluded only under a preregistered rule and remains in the exclusion log.')

d.add_heading('6.4 Architecture baselines',level=2)
para(d,'The primary baselines are a single tool-using agent, a three-stage sequential pipeline, parallel workers with an aggregator, and a manager–specialist–verifier hierarchy. Role prompts are minimal and public. The same backbone model is used across roles in the primary comparison; a heterogeneous-model extension is secondary. Total input/output token caps, tool-call caps, timeout, retriever access, and final validator are budget-matched. A secondary “natural budget” comparison allows each architecture to use its ordinary compute, clearly separating ecological performance from topology-isolation evidence.')
figure(d,'figure_04_topologies.png','Figure 7. Four architecture baselines. The primary comparison is budget-matched; natural-budget results are secondary.')

d.add_heading('6.5 Data splits and sample size',level=2)
para(d,'Tasks are divided into development, pilot, and locked evaluation splits. The pilot estimates clean success, within-task correlation, injection success, and variance for power simulations; pilot episodes are never included in confirmatory tests. The intended floor is 100 paired episode IDs per primary contrast, but this is not treated as a universal power guarantee. Final sample size is selected by simulation to achieve at least 80% power for the smallest practically important effect declared before the locked run—for binary task success, an absolute paired risk difference of 0.10 is the default planning target unless domain risk justifies a smaller threshold. If budget prevents adequate power, the study reports uncertainty and is labelled exploratory rather than manufacturing significance.')

# 7 experiments
d.add_heading('7. Experimental Design',level=1)
para(d,'The six experiments use paired task/seed blocks. Conditions are interleaved to reduce time-of-day, provider, and API-version confounding. The run manifest is generated before execution, cryptographically hashed, and never rewritten; failed launches, timeouts, invalid injections, and exclusions are retained with reason codes.')
figure(d,'figure_05_experiments.png','Figure 8. Six-experiment program and common controls.')

d.add_heading('7.1 Experiment 1 — Baseline multidimensional reliability (RQ1, H6)',level=2)
para(d,'E1 runs all four architectures on the same locked tasks with no injected fault across repeated seeds. Primary outcomes are task success, pass^k, pairwise end-state consistency, constraint-level success, tokens, tool calls, latency, and cost. H6 is supported only if a prespecified architecture ordering by task success disagrees materially with at least one other dimension or if systems with statistically indistinguishable task success differ meaningfully on repeatability, robustness in E2, or resource use. Rankings are not forced when confidence intervals overlap substantially.')

d.add_heading('7.2 Experiment 2 — Fault type, severity, and position (RQ2, H1)',level=2)
para(d,'E2 applies valid operators across early, middle, and late topological positions and calibrated severity levels. Each fault episode is paired with a null-intervention episode using the same task, seed block, architecture, model, and environment snapshot. Primary outcomes are EPR, DAF, amplification, depth, task-success degradation, and time to detection. H1 uses DAF and depth as primary metrics and includes reachable-set size as a covariate. Operator application rate and consumption rate are reported before propagation so silent injection failure cannot be misread as robustness.')

d.add_heading('7.3 Experiment 3 — Attribution under observability (RQ3, H2)',level=2)
para(d,'E3 constructs a locked set of failed episodes with gold agent and step labels. Each attribution method sees the same episode in three paired views: output only; partial trace containing inter-agent messages and tool names/status; and full trace containing prompts/inputs permitted by policy, messages, complete tool I/O, state deltas, timestamps, and dependency identifiers. Gold labels and task outcomes are hidden. Primary outcomes are joint exact accuracy and step localization distance; secondary outcomes are agent accuracy, top-k recall, mean reciprocal rank, abstention quality, and diagnostic cost. E3 tests evidence availability, not a change in the underlying failures.')

d.add_heading('7.4 Experiment 4 — Verification placement (RQ4, H3)',level=2)
para(d,'E4 compares no explicit verification, final-only verification, and local verification before critical handoffs. Verification rules are executable whenever possible: schema checks, constraint checks, unit tests, independent tool re-queries, or consistency checks. Primary outcomes are EPR and DAF; secondary outcomes include task success, false-rejection rate, false-acceptance rate, tokens, and latency. The local verifier is budgeted and versioned. To avoid giving it privileged gold answers, it receives only information available at its execution point.')

d.add_heading('7.5 Experiment 5 — Recovery strategies (RQ5, H4)',level=2)
para(d,'E5 starts from matched detected-fault episodes and compares: no recovery; same-context retry; retry with a concise error signal; rollback to the latest valid checkpoint; independent alternate-agent re-execution; and a bounded human gate for a high-risk subset. Recovery rate counts episodes that end in a correct state after a real, consumed fault. Residual propagation, repeated-failure similarity, added tokens, latency, tool calls, and cost are reported. The primary decision artifact is a Pareto frontier rather than a universal winner.')
figure(d,'figure_08_mitigation.png','Figure 9. Mitigation ladder used in Experiment 5 and the ablation. The figure specifies conditions; it does not show results.')

d.add_heading('7.6 Experiment 6 — Topology sensitivity (RQ6, H5)',level=2)
para(d,'E6 repeats selected clean and fault conditions across the four budget-matched architectures. Injection targets are mapped by functional role and normalized depth rather than agent name. For the parallel architecture, isolated worker faults and aggregator faults are separated; for the hierarchy, manager and leaf faults are separated. Primary analysis tests topology × fault-location interactions on DAF and task success. A secondary natural-budget analysis describes deployed behaviour but is not used to isolate topology causally.')

d.add_heading('7.7 Ablation study',level=2)
para(d,'The full MAS-RELIAB treatment uses verified injection, instrumented dependencies, local verification, checkpoints, full trace capture, and alternate-agent recovery. One component is removed at a time while holding the task/seed block fixed: no application check, no consumption check, no graph lineage, output-only evidence, no local verification, no checkpoint, and same-agent-only recovery. The ablation asks whether each component changes measurement validity or containment—not whether “MAS-RELIAB” as a label wins against an artificially weak baseline.')

# 8 metrics
d.add_heading('8. Metrics and Computation',level=1)
d.add_heading('8.1 Task outcome and repeatability',level=2)
equation(d,'TSR = N_successful / N_total','(9)')
equation(d,'pass^k = (1/|T|) Σ_{τ∈T} 𝟙[all k prespecified repeats of τ succeed]','(10)')
equation(d,'Consistency_k = (1/|T|) Σ_τ [2/(k(k−1))] Σ_{a<b} 𝟙[state_a ≡ state_b]','(11)')
para(d,'TSR includes all eligible launched episodes; timeouts and exhausted budgets are failures unless the task protocol specifies abstention as acceptable. pass^k is computed only at prespecified k values and requires a complete repeated-run block. Pairwise consistency uses task-defined end-state equivalence, not text identity.')

d.add_heading('8.2 Perturbation and fault robustness',level=2)
equation(d,'Δ_fault = TSR_clean − TSR_fault,     RR = N_correct_after_recovery / N_valid_consumed_faults','(12)')
para(d,'Robustness is reported as paired risk difference and risk ratio with confidence intervals. Fault tolerance is not inferred from a low observed error rate until injection validity and consumption are established. Results are stratified by operator and severity before any pooled estimate.')

d.add_heading('8.3 Propagation and containment',level=2)
para(d,'EPR, DAF, AF, depth, and edge transmission follow Equations (5)–(8). Additional metrics are failure containment score FCS=1−DAF, time to first valid detection, and lineage precision/recall when predicted propagation paths are compared with gold or replay-validated paths. For graph-size comparisons, both raw affected-node count and DAF are shown because normalization can hide large absolute impacts.')

d.add_heading('8.4 Attribution and detection',level=2)
equation(d,'Acc_agent = N_correct_agent / N_labelled,    Acc_step = N_correct_step / N_labelled','(13)')
equation(d,'MRR = (1/N) Σ_e 1/rank_e,    LD = (1/N) Σ_e dist_trace(t̂_e, t*_e)','(14)')
para(d,'Joint attribution requires both agent and step to be correct. A prediction outside the allowed evidence view or after the evaluation deadline is invalid. Detection reports precision, recall, F1, false alarms per episode, and time-to-detect. Abstentions are scored explicitly rather than silently removed.')

d.add_heading('8.5 Resource and trade-off metrics',level=2)
para(d,'Every episode records input, cached-input, reasoning (if exposed), and output tokens; model and tool calls; wall-clock and active compute latency; retries; checkpoint storage; and currency cost using a dated pricing snapshot. A system is Pareto-dominated when another condition is at least as good on all selected reliability dimensions and strictly better on at least one while costing no more. Incremental cost per additional recovered episode is reported with uncertainty. Ratios such as success per dollar are secondary because they can become unstable near zero cost.')

# 9 statistics
d.add_heading('9. Statistical Analysis Plan',level=1)
d.add_heading('9.1 Unit of inference and pairing',level=2)
para(d,'The task instance is the primary unit of inference. Repeated seeds and multiple faults within a task are nested observations, not independent task samples. Primary comparisons are paired by task, seed block, environment snapshot, and—where meaningful—fault instance. Clustered bootstrap resampling occurs at the task level. Results across task families are shown separately before pooled models.')

d.add_heading('9.2 Primary tests and models',level=2)
stat_rows=[
 ('Paired binary success','Exact McNemar test; paired risk difference and odds ratio','Hierarchical logistic mixed model with task random intercept.'),
 ('DAF / bounded fraction','Paired permutation or Wilcoxon signed-rank; median difference','Mixed beta/binomial model where assumptions hold; mass at 0/1 handled explicitly.'),
 ('Affected-node counts / depth','Paired permutation; Hodges–Lehmann shift','Negative-binomial or ordinal mixed model.'),
 ('Latency / tokens / cost','Paired bootstrap CI; Wilcoxon signed-rank','Log-linear mixed model after diagnostic checks.'),
 ('Attribution accuracy across views','Cochran Q followed by paired McNemar contrasts','Logistic mixed model with method and view effects.'),
 ('Rankings / trade-offs','Bootstrap rank stability and Pareto membership','No forced scalar omnibus score.')]
table(d,['Outcome','Primary analysis','Model-based secondary analysis'],stat_rows,widths=[1.55,2.4,2.55],small=8.0)

d.add_heading('9.3 Uncertainty, multiplicity, and effect size',level=2)
para(d,'All primary estimates receive 95% confidence intervals, preferably task-clustered BCa bootstrap intervals where computationally feasible. Confirmatory tests use two-sided α=0.05 even for directional hypotheses, and Holm correction within each research-question family. The manuscript reports exact p-values, not merely p<0.05. Effect sizes accompany tests: paired risk difference and matched odds ratio for binary outcomes; Hodges–Lehmann median shift or Cliff’s delta for skewed continuous outcomes; and standardized coefficients only when their assumptions are defensible. Practical-importance thresholds are preregistered separately from statistical significance.')

d.add_heading('9.4 Missing data, failures, and exclusions',level=2)
para(d,'Timeouts, budget exhaustion, provider errors, and parser failures are outcome categories and are not silently dropped. Injection failures are excluded only from propagation conditional on a valid consumed injection; they remain in the injection-reliability table. Planned missingness, such as unavailable internal prompts under a privacy policy, defines an observability condition. Any post-lock exclusion requires a reason code, count, sensitivity analysis with the episode retained where possible, and a signed amendment.')

d.add_heading('9.5 Reproducible analysis lock',level=2)
para(d,'The analysis script version, package lock, task manifest hash, run manifest hash, table specifications, hypothesis-to-metric mapping, and random seeds are frozen before condition labels are unblinded. A second script validates denominators and reconstructs a random sample of episode metrics directly from raw traces. The final paper records deviations from this plan.')

# 10 Results shell
d.add_heading('10. Results — Locked Reporting Shell',level=1)
callout(d,'Results status.','No empirical runs were provided. Every em dash in this section is intentional. Do not replace it with an illustrative number. Populate only from the frozen analysis pipeline and retain contradictory or null findings.',RED)

d.add_heading('10.1 Run accounting and data quality',level=2)
caption(d,'Table 2. Episode accounting. Denominators must reconcile with the signed run manifest.')
table(d,['Stage','Planned','Launched','Completed','Timed out','Invalid injection','Excluded','Analysed'],[
 ('Clean baseline','—','—','—','—','n/a','—','—'),('Fault injection','—','—','—','—','—','—','—'),('Attribution set','—','—','—','—','n/a','—','—'),('Verification','—','—','—','—','—','—','—'),('Recovery','—','—','—','—','—','—','—'),('Topology','—','—','—','—','—','—','—')],small=8.0)
para(d,'Required narrative: report task counts by family, model and environment versions, execution dates, provider incidents, injection application and consumption rates, exclusion reasons, and any preregistration deviations. No result interpretation begins until this reconciliation is complete.')

d.add_heading('10.2 E1: Baseline reliability',level=2)
caption(d,'Table 3. Baseline reliability by architecture; values are estimate [95% CI].')
table(d,['Architecture','TSR','pass^k','Consistency','Tokens','Latency','Cost','H6 note'],[
 ('Single agent','—','—','—','—','—','—','—'),('Sequential','—','—','—','—','—','—','—'),('Parallel + aggregator','—','—','—','—','—','—','—'),('Hierarchical','—','—','—','—','—','—','—')],small=8.0)
para(d,'Decision text to complete: “Architecture rankings were [consistent/inconsistent] across task success and the preregistered reliability dimensions. H6 was [supported/not supported/inconclusive] because —.”')

d.add_heading('10.3 E2: Propagation by fault and position',level=2)
caption(d,'Table 4. Conditional propagation after verified application and consumption.')
table(d,['Fault family','Position','N valid','EPR','DAF','Amplification','Depth','Δ task success'],[
 ('Agent-local','Early','—','—','—','—','—','—'),('Agent-local','Middle','—','—','—','—','—','—'),('Agent-local','Late','—','—','—','—','—','—'),('Handoff','Early / Middle / Late','—','—','—','—','—','—'),('Tool/environment','Early / Middle / Late','—','—','—','—','—','—'),('Coordination','Early / Middle / Late','—','—','—','—','—','—')],small=7.7)
para(d,'Decision text to complete: “After reachable-set normalization and multiplicity correction, early injections had — DAF and — depth than late injections. H1 was —. The strongest operator-specific effect was —; injection failure or non-consumption accounted for —.”')
caption(d,'Figure 10 placeholder. Failure propagation by normalized topological position. Generate from analysis/fig_e2_position.*; include points, task-clustered 95% CIs, valid-injection counts, and no smoothed line unless prespecified.')

d.add_heading('10.4 E3: Attribution and observability',level=2)
caption(d,'Table 5. Paired attribution performance on identical failed episodes.')
table(d,['Evidence view','Joint exact','Agent exact','Step exact','Top-k recall','MRR','Localization distance','Diagnostic cost'],[
 ('Output only','—','—','—','—','—','—','—'),('Partial trace','—','—','—','—','—','—','—'),('Full trace','—','—','—','—','—','—','—')],small=7.7)
para(d,'Decision text to complete: “On the same locked episodes, full observability changed joint attribution by — relative to output-only evidence. H2 was —. Gains/losses were concentrated in —, while privacy or logging constraints affected —.”')
caption(d,'Figure 11 placeholder. Agent- and step-level attribution by evidence view with paired confidence intervals and abstention rate.')

d.add_heading('10.5 E4: Verification',level=2)
caption(d,'Table 6. Verification placement and containment.')
table(d,['Regime','TSR','EPR','DAF','False accept','False reject','Added tokens','Added latency'],[
 ('No explicit verification','—','—','—','—','—','—','—'),('Final-only','—','—','—','—','—','—','—'),('Local before handoff','—','—','—','—','—','—','—')],small=8.0)
para(d,'Decision text to complete: “Local verification changed EPR by — and DAF by — relative to final-only checking, with — added resource use. H3 was —. False interventions were —.”')

d.add_heading('10.6 E5: Recovery',level=2)
caption(d,'Table 7. Recovery effectiveness and resource trade-off.')
table(d,['Recovery strategy','Recovery rate','Residual DAF','Repeated-failure similarity','Added tokens','Added latency','Added cost','Pareto status'],[
 ('No recovery','—','—','—','—','—','—','—'),('Same-context retry','—','—','—','—','—','—','—'),('Retry + error signal','—','—','—','—','—','—','—'),('Rollback','—','—','—','—','—','—','—'),('Alternate agent','—','—','—','—','—','—','—'),('Human gate subset','—','—','—','—','—','—','—')],small=7.4)
para(d,'Decision text to complete: “The Pareto frontier contained —. Relative to same-context retry, rollback changed recovery by — and alternate-agent recovery changed it by —. H4 was —; the incremental cost per additional recovered episode was —.”')
caption(d,'Figure 12 placeholder. Reliability–cost Pareto frontier with uncertainty and dominated conditions labelled.')

d.add_heading('10.7 E6: Topology sensitivity',level=2)
caption(d,'Table 8. Budget-matched topology × fault-location interaction.')
table(d,['Topology','Fault location','TSR','EPR','DAF','Depth','Interaction estimate','H5 interpretation'],[
 ('Single','Agent / tool','—','—','—','—','—','—'),('Sequential','Early / late','—','—','—','—','—','—'),('Parallel','Worker / aggregator','—','—','—','—','—','—'),('Hierarchical','Leaf / manager','—','—','—','—','—','—')],small=7.7)
para(d,'Decision text to complete: “The topology × location interaction was —. Isolated worker faults were —, aggregator faults were —, leaf faults were —, and manager faults were — after budget matching. H5 was —.”')
caption(d,'Figure 13 placeholder. Heatmap of DAF by topology and normalized fault position, annotated with valid-injection counts.')

d.add_heading('10.8 Ablation and robustness checks',level=2)
caption(d,'Table 9. One-component-at-a-time ablation; report change from the full condition.')
table(d,['Ablation','Δ TSR','Δ EPR','Δ DAF','Δ attribution','Δ cost','Validity consequence'],[
 ('No applied check','—','—','—','—','—','—'),('No consumption check','—','—','—','—','—','—'),('No graph lineage','—','—','—','—','—','—'),('Output-only evidence','—','—','—','—','—','—'),('No local verification','—','—','—','—','—','—'),('No checkpoint','—','—','—','—','—','—'),('Same-agent-only recovery','—','—','—','—','—','—')],small=7.7)
para(d,'Robustness checks to report: alternate success equivalence, per-family estimates, exclusion-retention sensitivity, natural-budget comparison, provider-time fixed effects, alternate graph-edge rules, multiple-root attribution scoring, and weight sensitivity if any composite is explored.')

# 11 Discussion
d.add_heading('11. Discussion Framework',level=1)
para(d,'This section states interpretation rules, not empirical findings. It must be rewritten after the locked Results are available.')
d.add_heading('11.1 Interpreting RQ1 and H6',level=2)
para(d,'If task-success rankings diverge from pass^k, fault degradation, propagation containment, or cost, the defensible conclusion is that single-run success masked a dimension relevant to deployment. If all dimensions agree, H6 is not supported; MAS-RELIAB may still provide diagnostic detail, but rank reversal cannot be claimed. A system that is accurate and expensive is not “unreliable” solely because of cost; reliability and resources remain distinct axes.')

d.add_heading('11.2 Interpreting position and propagation',level=2)
para(d,'A larger raw cascade from early positions may simply reflect more reachable nodes. H1 therefore depends on normalized DAF and depth with reachable-set adjustment. If early and late effects become similar after normalization, the correct interpretation is opportunity-driven spread, not greater intrinsic early-node transmissibility. Operator-specific heterogeneity should not be erased by a pooled average.')

d.add_heading('11.3 Interpreting observability and attribution',level=2)
para(d,'Better attribution with full traces would support instrumentation investment, but it would not prove that every internal reasoning token should be logged. Operational systems must balance diagnostic value with privacy, security, intellectual-property, and storage constraints. If full traces do not help, possible explanations include weak attribution methods, noisy logs, inadequate gold labels, or genuinely ambiguous multiple causes; the study should test these alternatives rather than asserting that observability is irrelevant.')

d.add_heading('11.4 Interpreting mitigation',level=2)
para(d,'Verification and recovery can reduce propagation while increasing cost, latency, or false rejection. The appropriate policy is risk-tiered: inexpensive local checks for common deterministic constraints, checkpoint/alternate-agent recovery for consequential handoffs, and human escalation only where residual risk justifies delay. No mitigation is universally best without a stated utility or safety threshold.')

d.add_heading('11.5 Engineering implications',level=2)
bullets(d,[
 'Instrument causal provenance at message, tool, and state boundaries rather than storing final outputs only.',
 'Validate that a chaos/fault operator actually reached the intended consumer before calling the system tolerant.',
 'Place checks before high-centrality or high-consequence handoffs, not only at final output.',
 'Keep retries bounded and compare them with rollback or independent re-execution; retries can repeat the same correlated error.',
 'Use reliability–cost Pareto frontiers and risk thresholds instead of a hidden weighted score.',
 'Treat topology as an interaction with fault location and aggregation policy, not as a global label such as “parallel is robust.”'
])

# 12 reproducibility
d.add_heading('12. Reproducibility, Governance, and Ethics',level=1)
d.add_heading('12.1 Required run record',level=2)
para(d,'For every episode, record: task and dataset version; gold-state hash; architecture/topology; agent roles; full system and role prompts or governed hashes; model provider, exact model/version, parameters, decoding seed and endpoint region; tool versions and snapshots; orchestration and MAS-RELIAB commit IDs; dependency lock; operating system/container digest; hardware where locally relevant; environment state; fault operator, target, position, severity, seed, patch and checks; timeout and token/tool budgets; all messages and tool I/O allowed by policy; validator output; start/end timestamps; tokens, calls, latency and cost; and inclusion/exclusion status.')

d.add_heading('12.2 Artifact release',level=2)
para(d,'Release the preregistration, benchmark schema, public task split where licensing permits, injection operators, minimal role prompts, environment snapshots or builders, anonymized traces, graph-construction code, validators, analysis scripts, table/figure specifications, run manifest, package lock, and model-card-style limitations. If proprietary APIs or data prevent full release, provide hashes, synthetic substitutes, replay fixtures, and a precise restricted-access statement.')

d.add_heading('12.3 Ethical and security controls',level=2)
para(d,'Fault injection is confined to sandboxed or simulated environments. Operators must not target live third-party services, real customer records, or production credentials. Prompt and trace logging follows least-privilege access, retention limits, redaction, and documented consent/terms. Adversarial payloads that could facilitate misuse are summarized in public artifacts and shared in full only when risk assessment permits. Human evaluators receive guidance for potentially harmful content and may abstain without penalty.')

d.add_heading('12.4 Reproducibility checklist',level=2)
check_items=['Preregistration timestamp and immutable hash','Task manifest and split hashes','Model/provider/version and decoding parameters','Prompt and role files or governed hashes','Topology and routing configuration','Fault operator patch, target, severity, seed, applied/consumed checks','Environment/container/tool versions','Hardware and region where relevant','Run schedule and provider incident log','Raw trace and state snapshots','Validator version and constraint outputs','Token, tool-call, latency and pricing snapshot','Exclusion and missing-data log','Analysis environment lock and seeds','Exact table/figure generation command','Deviations and sensitivity analyses']
for item in check_items: para(d,'☐  '+item)

# 13 limits
d.add_heading('13. Limitations',level=1)
bullets(d,[
 ('Pre-experimental status. ','The current document specifies the protocol but contains no empirical validation. Feasibility, runtime, annotation burden, and metric behaviour remain to be demonstrated.'),
 ('Benchmark representativeness. ','Five task families cannot represent every deployment, especially embodied, continuous, creative, or human-centred work.'),
 ('Fault realism. ','Controlled operators improve causal identification but may not reproduce the joint distribution of naturally occurring failures.'),
 ('Graph completeness. ','Instrumented dependencies omit latent model associations; inferred semantic edges may be uncertain.'),
 ('Counterfactual instability. ','Replaying a stochastic system after one repaired event can produce a different trajectory for reasons unrelated to the repair.'),
 ('Provider drift. ','Hosted models, rate limits, safety policies, and pricing can change despite interleaving and version records.'),
 ('Ground-truth ambiguity. ','Some failures have multiple sufficient causes or no unique earliest decisive step.'),
 ('Observability constraints. ','Full trace capture may be impossible or undesirable in privacy-sensitive and proprietary environments.'),
 ('Budget matching. ','Equal token/tool budgets improve internal validity but may disadvantage architectures designed for greater parallel compute.'),
 ('Statistical power and multiplicity. ','The factorial space is large; some interactions may remain underpowered even with substantial runs.'),
 ('Composite avoidance does not remove value judgments. ','Selecting which axes and practical thresholds matter still requires a deployment-specific risk decision.'),
 ('External comparison. ','Related work is rapidly evolving; scope-table cells and novelty language must be reverified at submission time.')
])

# 14 threats
d.add_heading('14. Threats to Validity',level=1)
validity_rows=[
 ('Construct validity','Metrics may not fully represent “reliability”; validators may reward surface correctness; fault doses may not reflect meaningful severity.','Use component metrics, state constraints, multiple operators, pilot calibration, and transparent definitions.'),
 ('Internal validity','Model drift, scheduler load, different reachable sets, injection non-consumption, or budget imbalance may confound effects.','Pair and interleave runs, pin snapshots, verify injection, normalize opportunity, and budget-match primary contrasts.'),
 ('External validity','Synthetic tasks, selected models, and fixed team sizes may not generalize.','Use diverse task families/models, report strata, release adapters, and avoid universal claims.'),
 ('Conclusion validity','Nested repeats treated as independent, multiple testing, low power, or selective exclusions could distort inference.','Task-level clustering, mixed models, Holm correction, power simulation, effect sizes, CIs, and immutable manifests.'),
 ('Reproducibility validity','Undocumented prompts, endpoint versions, tool state, or pricing can prevent replication.','Full run records, environment digests, manifests, lockfiles, and replay fixtures.'),
 ('Evaluation bias','LLM judges or annotators may prefer certain styles/models.','Deterministic primary validators, blinded labels, adjudication, agreement statistics, and judge sensitivity checks.')]
table(d,['Validity type','Threat','Mitigation'],validity_rows,widths=[1.25,2.75,2.55],small=8.0)

# 15 conclusion
d.add_heading('15. Conclusion',level=1)
para(d,'MAS-RELIAB has been specified as a results-ready methodology for evaluating reliability and failure propagation in multi-agent AI systems. It represents executions as instrumented dependency graphs, requires documented and verified fault interventions, reports reliability as interpretable components, measures propagation and attribution under explicit observability conditions, compares verification and recovery strategies, and controls topology and resource use. AI Command Center serves as the execution and observability substrate; MAS-RELIAB remains the benchmark, fault-injection, metrics, and analysis layer.')
para(d,'No empirical conclusion is yet warranted. The six experiments, H1–H6, result tables, and decision rules are prepared for execution, but their outcomes are unknown. The final conclusion must state the observed effects with confidence intervals and effect sizes, disclose null or contradictory results, and revise claims accordingly. At present, the defensible contribution is the integrated experimental protocol—not demonstrated superiority and not a priority claim for graph-based cascade analysis or mitigation.')

# References
d.add_heading('References',level=1)
refs=[
 '[1] M. Cemri et al., “Why Do Multi-Agent LLM Systems Fail?” arXiv:2503.13657, 2025. https://arxiv.org/abs/2503.13657',
 '[2] K. Zhu et al., “MultiAgentBench: Evaluating the Collaboration and Competition of LLM Agents,” arXiv:2503.01935, 2025. https://arxiv.org/abs/2503.01935',
 '[3] S. Yao, N. Shinn, P. Razavi, and K. Narasimhan, “τ-bench: A Benchmark for Tool-Agent-User Interaction in Real-World Domains,” arXiv:2406.12045, 2024. https://arxiv.org/abs/2406.12045',
 '[4] A. Gupta, “ReliabilityBench: Evaluating LLM Agent Reliability Under Production-Like Stress Conditions,” arXiv:2601.06112, 2026. https://arxiv.org/abs/2601.06112',
 '[5] M. Chen et al., “Seeing the Whole Elephant: A Benchmark for Failure Attribution in LLM-based Multi-Agent Systems,” ACL 2026; arXiv:2604.22708. https://arxiv.org/abs/2604.22708',
 '[6] G. Zhang, J. Wang, J. Chen, W. Zhou, K. Wang, and S. Yan, “AgenTracer: Who Is Inducing Failure in the LLM Agentic Systems?” arXiv:2509.03312, 2025. https://arxiv.org/abs/2509.03312',
 '[7] J. Li, E. Yilmaz, B. Chen, and T. Le, “Towards Self-Improving Error Diagnosis in Multi-Agent Systems,” Findings of ACL 2026, pp. 2063–2077. https://aclanthology.org/2026.findings-acl.98/',
 '[8] B. Lin et al., “AgentAsk: Multi-Agent Systems Need to Ask,” ACL 2026, pp. 28055–28077. https://aclanthology.org/2026.acl-long.1294/',
 '[9] Y. Jiang et al., “RiskLab: A Controlled Toolkit for Probing Emergent Risks in LLM-Based Multi-Agent Systems,” ACL 2026 System Demonstrations, pp. 167–177. https://aclanthology.org/2026.acl-demo.17/',
 '[10] I. Kavathekar, H. Jain, A. Rathod, P. Kumaraguru, and T. Ganu, “TAMAS: Benchmarking Adversarial Risks in Multi-Agent LLM Systems,” arXiv:2511.05269, 2025. https://arxiv.org/abs/2511.05269',
 '[11] Y. Zhang et al., “Silo-Bench: A Scalable Environment for Evaluating Distributed Coordination in Multi-Agent LLM Systems,” ACL 2026; arXiv:2603.01045. https://arxiv.org/abs/2603.01045',
 '[12] Y. Xie et al., “From Spark to Fire: Modeling and Mitigating Error Cascades in LLM-Based Multi-Agent Collaboration,” arXiv:2603.04474, 2026. https://arxiv.org/abs/2603.04474',
 '[13] A. Yehudai et al., “Survey on Evaluation of LLM-based Agents,” arXiv:2503.16416, 2025. https://arxiv.org/abs/2503.16416',
 '[14] M. Mohammadi, Y. Li, J. Lo, and W. Yip, “Evaluation and Benchmarking of LLM Agents: A Survey,” Proc. 31st ACM SIGKDD, 2025; arXiv:2507.21504. https://arxiv.org/abs/2507.21504',
 '[15] P. Zhu et al., “A Unified Framework for the Evaluation of LLM Agentic Capabilities,” arXiv:2605.27898, 2026. https://arxiv.org/abs/2605.27898',
 '[16] S. Yao et al., “ReAct: Synergizing Reasoning and Acting in Language Models,” International Conference on Learning Representations, 2023. https://arxiv.org/abs/2210.03629',
 '[17] N. Shinn et al., “Reflexion: Language Agents with Verbal Reinforcement Learning,” Advances in Neural Information Processing Systems, vol. 36, 2023. https://arxiv.org/abs/2303.11366',
 '[18] Q. Wu et al., “AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation,” arXiv:2308.08155, 2023. https://arxiv.org/abs/2308.08155',
 '[19] C. E. Jimenez et al., “SWE-bench: Can Language Models Resolve Real-World GitHub Issues?” International Conference on Learning Representations, 2024. https://arxiv.org/abs/2310.06770',
 '[20] G. Mialon et al., “GAIA: A Benchmark for General AI Assistants,” International Conference on Learning Representations, 2024. https://arxiv.org/abs/2311.12983',
 '[21] X. Liu et al., “AgentBench: Evaluating LLMs as Agents,” International Conference on Learning Representations, 2024. https://arxiv.org/abs/2308.03688',
 '[22] National Institute of Standards and Technology, “Artificial Intelligence Risk Management Framework (AI RMF 1.0),” NIST AI 100-1, 2023. https://doi.org/10.6028/NIST.AI.100-1',
 '[23] B. Efron and R. J. Tibshirani, An Introduction to the Bootstrap. Chapman & Hall/CRC, 1993.',
 '[24] S. Holm, “A Simple Sequentially Rejective Multiple Test Procedure,” Scandinavian Journal of Statistics, vol. 6, no. 2, pp. 65–70, 1979.',
 '[25] N. Cliff, Ordinal Methods for Behavioral Data Analysis. Lawrence Erlbaum Associates, 1996.'
]
for r in refs:
    p=d.add_paragraph(r); p.paragraph_format.left_indent=Inches(.22); p.paragraph_format.first_line_indent=Inches(-.22); p.paragraph_format.space_after=Pt(3)

# Appendices
page_break(d)
d.add_heading('Appendix A. Hypothesis–Metric–Decision Matrix',level=1)
h_rows=[
 ('H1','DAF; depth','Early vs late valid consumed injections','Holm-adjusted paired contrast and practically meaningful direction','Reject/retain/inconclusive; explain opportunity adjustment.'),
 ('H2','Joint attribution; localization distance','Full vs partial vs output-only on same failures','Paired evidence-view effect with 95% CI','Report privacy/logging caveat.'),
 ('H3','EPR; DAF','Local vs final-only vs none','Paired reduction without unacceptable false rejection','Show reliability and cost jointly.'),
 ('H4','Recovery rate; cost','Rollback/alternate vs same-context retry','Paired recovery gain and Pareto status','No universal winner claim.'),
 ('H5','DAF; TSR interaction','Topology × functional fault location','Prespecified interaction estimate','Separate worker/aggregator and leaf/manager.'),
 ('H6','Rank discordance','TSR ranking vs other dimensions','Prespecified discordance or masked weakness','No support if dimensions agree.')]
table(d,['Hyp.','Primary metric','Primary contrast','Decision evidence','Required wording'],h_rows,widths=[.5,1.3,1.8,2.0,1.55],small=7.8)

d.add_heading('Appendix B. Machine-Readable Schemas',level=1)
d.add_heading('B.1 Task schema (abridged)',level=2)
code='''{
  "task_id": "string", "task_version": "semver", "family": "enum",
  "instruction": "string", "initial_state_digest": "sha256",
  "tools": [{"name":"string","version":"string"}],
  "gold": {"required_constraints": [], "equivalence_rules": []},
  "budgets": {"tokens": null, "tool_calls": null, "seconds": null},
  "injection_points": [{"event_selector":"string","allowed_operators":[]}],
  "split": "dev|pilot|locked", "provenance": {}, "license": "string"
}'''
p=d.add_paragraph(); r=p.add_run(code); r.font.name='DejaVu Sans Mono'; r.font.size=Pt(8); shade_dummy=None

d.add_heading('B.2 Trace event schema (abridged)',level=2)
code='''{
  "run_id":"uuid", "event_id":"uuid", "parent_event_ids":[],
  "actor_id":"string", "role":"string", "event_type":"enum",
  "timestamp_utc":"RFC3339", "model_or_tool_version":"string",
  "input_hashes":[], "output_hash":"sha256", "state_delta":{},
  "tokens":{}, "latency_ms":null, "cost":null,
  "fault_lineage_ids":[], "visibility":"output|partial|full",
  "validator_annotations":[]
}'''
p=d.add_paragraph(); r=p.add_run(code); r.font.name='DejaVu Sans Mono'; r.font.size=Pt(8)

d.add_heading('B.3 Fault intervention schema (abridged)',level=2)
code='''{
  "intervention_id":"uuid", "operator_version":"string", "family":"enum",
  "target_selector":"string", "trigger":"string", "position":{},
  "severity":{}, "seed":null, "payload_digest":"sha256",
  "precondition":{}, "applied_check":{}, "consumed_check":{},
  "before_hash":"sha256", "after_hash":"sha256", "status":"enum"
}'''
p=d.add_paragraph(); r=p.add_run(code); r.font.name='DejaVu Sans Mono'; r.font.size=Pt(8)

d.add_heading('Appendix C. Core Pseudocode',level=1)
d.add_heading('C.1 Controlled episode',level=2)
code='''for block in signed_run_manifest:
    restore(block.environment_snapshot)
    configure(block.model, block.topology, block.prompts, block.budgets)
    if block.intervention is not null:
        assert precondition(block.intervention)
        patch = apply(block.intervention)
        log(before_hash, patch, after_hash)
    run_id = AI_Command_Center.execute_episode(block.config)
    trace = export_and_validate_trace(run_id)
    applied  = verify_application(trace, block.intervention)
    consumed = verify_consumption(trace, block.intervention)
    outcome  = deterministic_validator(final_state(run_id), block.gold)
    graph    = build_instrumented_dependency_graph(trace)
    store_immutable(run_id, applied, consumed, outcome, graph, resources(trace))'''
p=d.add_paragraph(); r=p.add_run(code); r.font.name='DejaVu Sans Mono'; r.font.size=Pt(7.8)

d.add_heading('C.2 Propagation labelling',level=2)
code='''origin = gold_origin(intervention_log, replay_evidence)
reachable = descendants(graph, origin)
affected = set()
for node in reachable ordered by time:
    if consumes_fault_lineage(node, origin) and task_relevant_error(node):
        if not explained_by_logged_independent_cause(node):
            affected.add(node)
EPR = int(len(affected) > 0)
DAF = len(affected) / len(reachable) if reachable else 0
AF = len(affected) / max(1, len(initially_corrupted_nodes))
Depth = max(shortest_path(graph, origin, v) for v in affected) if affected else 0'''
p=d.add_paragraph(); r=p.add_run(code); r.font.name='DejaVu Sans Mono'; r.font.size=Pt(7.8)

d.add_heading('Appendix D. Annotation Codebook',level=1)
ann_rows=[
 ('Origin','Earliest event that introduces a task-relevant incorrect state into the executed lineage.','Do not label a later agent that faithfully transmits an already incorrect upstream value.'),
 ('Decisive step','A repair at this event removes the labelled failure lineage or restores a feasible path under controlled replay.','If several steps are independently sufficient, use a set-valued label.'),
 ('Propagation','A downstream node consumes the lineage and newly exhibits a validated related error.','Temporal succession alone is insufficient.'),
 ('Containment','The lineage is detected, corrected, quarantined, or ignored before causing a downstream task-relevant error.','A branch with no reachable consumer is not evidence of active containment.'),
 ('Independent failure','A downstream error has a logged cause unrelated to the injected lineage.','Exclude it from A_i but retain it in the overall failure taxonomy.'),
 ('Ambiguous','Evidence cannot distinguish two or more plausible origins.','Allow abstention/set labels; do not force false precision.')]
table(d,['Label','Operational definition','Boundary rule'],ann_rows,widths=[1.1,3.0,2.2],small=8.0)

d.add_heading('Appendix E. Analysis and Reporting Checklist',level=1)
numbered(d,[
 'Reconcile every table denominator with the run manifest and report injection application/consumption separately.',
 'Show per-task-family and per-operator estimates before pooled results.',
 'Use task-clustered confidence intervals and paired tests for paired designs.',
 'Report effect size, exact p-value, adjusted p-value, and practical threshold for each confirmatory contrast.',
 'Retain timeouts and infrastructure failures as declared outcome categories.',
 'State whether a result is confirmatory, secondary, or exploratory.',
 'Never convert an unmeasured field into zero; use missing/unknown and explain.',
 'Generate plots only from analysis outputs; no hand-entered points.',
 'If a hypothesis is contradicted, report the observed direction and revise the discussion.',
 'Rewrite the abstract and conclusion only after all tables pass independent validation.'
])

# Viva
page_break(d)
d.add_heading('Viva Defence Pack',level=1)
d.add_heading('V1. Two-minute opening statement',level=2)
para(d,'Good morning. My name is Dev Parth, and my work is titled “A Framework for Evaluating Reliability and Failure Propagation in Multi-Agent AI Systems,” or MAS-RELIAB. The problem is that multi-agent systems are usually judged by whether the final task succeeded, but that single score cannot tell us whether success is repeatable, where a failure began, how it moved through the team, whether it could be correctly attributed, whether a mitigation contained it, or what that reliability cost.')
para(d,'MAS-RELIAB addresses this as an experimental methodology. It represents each execution as a directed, time-aware dependency graph linked to a complete trace and an objectively checked final state. It injects controlled, documented faults and verifies both that the fault was applied and that the intended component consumed it. It then reports task success, repeated-run consistency, fault degradation, propagation rate and depth, attribution accuracy, recovery, tokens, latency, and cost as separate dimensions rather than hiding them in one arbitrary score.')
para(d,'The protocol compares single-agent, sequential, parallel, and hierarchical systems. Six experiments test baseline reliability, fault position, observability for attribution, local versus final verification, recovery strategies, and topology sensitivity. AI Command Center is the execution and observability infrastructure; MAS-RELIAB provides the benchmark, fault operators, metrics, and analysis. My novelty claim is deliberately limited: prior work already studies taxonomies, fault injection, attribution, topologies, and graph-based error cascades. The contribution is the unified and auditable linkage of these components in one paired protocol. Because experiments have not yet been run, I do not claim empirical improvement. My current conclusion is that the methodology is specified and results-ready; all empirical hypotheses remain open.')

d.add_heading('V2. Ten-minute defence script',level=2)
script_rows=[
 ('0:00–0:40','Title and problem','A local error can become a coherent system failure through dependencies. Final success alone is not diagnostic.'),
 ('0:40–1:30','Gap and related work','Acknowledge MAST, τ-bench, ReliabilityBench, TraceElephant, AgentAsk, RiskLab, AgenTracer, and From Spark to Fire. State the integration claim only.'),
 ('1:30–2:10','Aim, RQs, hypotheses','Show RQ1–RQ6 and explain that hypotheses are preregistered and falsifiable.'),
 ('2:10–3:05','Architecture boundary','AI Command Center executes and logs; MAS-RELIAB defines tasks, injection, graph, metrics, and analysis.'),
 ('3:05–4:00','Formal model','Explain episode tuple, directed dependencies, valid consumed intervention, and end-state success.'),
 ('4:00–5:00','Fault taxonomy','Show typed operators, severity, position, and applied/consumed checks. Emphasize reproducibility.'),
 ('5:00–6:00','Metrics','Explain TSR, pass^k, EPR, DAF, depth, attribution, recovery, and cost. Reject arbitrary composite scoring.'),
 ('6:00–7:25','Six experiments','Walk through E1–E6 and identify each paired primary contrast.'),
 ('7:25–8:15','Statistics','Task is the unit; paired tests, mixed models, clustered 95% CIs, Holm correction, effect sizes, power simulation.'),
 ('8:15–9:00','Validity and ethics','Fault realism, graph incompleteness, provider drift, privacy, sandboxing, and budget matching.'),
 ('9:00–9:40','Results status','All values are pending. Explain why blanks demonstrate integrity rather than incompleteness.'),
 ('9:40–10:00','Close','MAS-RELIAB is a reproducible linkage of reliability, propagation, attribution, mitigation, topology, and cost; validation is the next step.')]
table(d,['Time','Slide focus','What to say'],script_rows,widths=[1.0,1.45,4.0],small=8.0)

d.add_heading('V3. Slide map',level=2)
slide_rows=[('1–3','Title, integrity status, motivating cascade'),('4–6','Evaluation gap, novelty boundary, related-work map'),('7–9','Aim/RQs/hypotheses and contributions'),('10–13','Layered architecture, graph formalism, taxonomy, topologies'),('14–17','Task design, reliability metrics, propagation, attribution views'),('18–21','Six experiments, verification, recovery, statistics'),('22–24','Results shell, reproducibility, limitations'),('25–27','Contribution boundary, next steps, conclusion'),('Appendix','Metric definitions, hypothesis matrix, examiner backup')]
table(d,['Slides','Purpose'],slide_rows,widths=[1.2,5.2],small=8.3)

d.add_heading('V4. Likely examiner questions and model answers',level=2)
qa=[
 ('1. What is the central research problem?','How to evaluate not only whether a multi-agent system succeeds, but whether success is repeatable and robust, how failures propagate, whether they can be attributed and contained, and what resources are required.'),
 ('2. What is your main contribution?','A unified, auditable experimental protocol linking multidimensional reliability, verified fault injection, propagation, observability-conditioned attribution, mitigation, topology, and cost in paired episodes.'),
 ('3. Are you claiming the first multi-agent reliability framework?','No. That claim would be indefensible. I explicitly position MAS-RELIAB as an integration and methodology contribution.'),
 ('4. How is this different from From Spark to Fire?','That work already models directed dependency graphs, atomic error cascades, topology effects, and a governance layer. MAS-RELIAB adds a broader controlled protocol tying cascade metrics to repeated reliability, evidence-view attribution, multiple verification/recovery baselines, state-based tasks, budget matching, and cost.'),
 ('5. How is this different from ReliabilityBench?','ReliabilityBench focuses on repeated consistency, perturbation, and tool/API faults for agents. MAS-RELIAB keeps those principles but makes multi-agent dependencies, propagation, attribution, topology, and mitigation central.'),
 ('6. How is this different from TraceElephant or AgenTracer?','They focus strongly on attribution and trace evidence. MAS-RELIAB treats attribution as one linked experiment and measures how evidence view, propagation, mitigation, outcome, and cost interact.'),
 ('7. Why use a directed graph?','Because information and control flow have direction. The graph makes reachability, lineage, topological position, transmission opportunities, and propagation depth explicit.'),
 ('8. Are graph edges causal?','Instrumented consumption edges show dependency opportunity, not automatically causal responsibility. Propagation labels additionally require temporal order, lineage evidence, and a task-relevant error; stronger causal claims use paired interventions or replay.'),
 ('9. Why not use only final task success?','It hides inconsistency, fault sensitivity, cascading impact, attribution, recovery, and resource overhead. Two systems with equal success may have very different operational risk.'),
 ('10. Why not combine everything into one reliability score?','The dimensions have different units and deployment values. Arbitrary weights can reverse rankings. I report components first and use Pareto analysis; any composite is exploratory with weight sensitivity.'),
 ('11. What is EPR?','The proportion of valid, consumed fault injections that cause at least one validated downstream failure.'),
 ('12. What is DAF?','The fraction of reachable downstream nodes that become affected. It normalizes for the propagation opportunity available from a position.'),
 ('13. Can EPR and DAF disagree?','Yes. A condition can propagate frequently but affect few nodes, or propagate rarely but cause large cascades when it does. Both are needed.'),
 ('14. How do you avoid making early faults look worse simply because they have more downstream nodes?','Use reachable-set-normalized DAF, report raw counts, include reachable-set size, and compare normalized topological positions.'),
 ('15. What makes a fault injection valid?','The operator’s precondition holds, the intended artifact changed as documented, and the target or downstream component actually consumed the altered artifact.'),
 ('16. Why distinguish application from consumption?','A correctly patched message that no agent reads cannot test tolerance. Without consumption checks, failed injections may be mistaken for robust behaviour.'),
 ('17. How is severity defined?','By operator-specific, measurable dose—such as fields altered, information removed, or calls failed—calibrated in the pilot and frozen before evaluation.'),
 ('18. How do you obtain ground truth?','Primary task truth comes from end-state constraints, tests, solvers, or closed-corpus fact keys. Fault origins come from intervention logs and replay; natural failures use blinded double annotation and adjudication.'),
 ('19. Why avoid an LLM judge as primary ground truth?','It can introduce model/style bias and nondeterminism. It may be a secondary diagnostic, but core tasks use objectively checkable state or constraints.'),
 ('20. What if a failure has multiple causes?','Use set-valued origin labels and set precision/recall. The framework permits ambiguity and abstention rather than forcing a false single root.'),
 ('21. What is the earliest decisive step?','The earliest event whose controlled repair removes the labelled lineage or restores a feasible path to success with the prior prefix fixed.'),
 ('22. Is counterfactual replay reliable with stochastic models?','Not perfectly. Replay instability is a limitation; use fixed seeds where supported, multiple replays, environment snapshots, and report repair-effect uncertainty.'),
 ('23. Why compare observability views on the same failures?','Otherwise a difference in attribution could be caused by a different episode mix rather than evidence availability. Pairing isolates the view effect.'),
 ('24. Does full observability mean recording private chain-of-thought?','No. Full means all legally and operationally available inputs, messages, tool I/O, state changes, timing, and provenance. Hidden reasoning is not assumed or required.'),
 ('25. Why local verification?','It can stop an error before the next dependency edge. Final-only checks may detect failure after the cascade has already consumed resources or changed state.'),
 ('26. Could local verification harm performance?','Yes. It may reject valid work, add latency, or create new failures. That is why false rejection and resource cost are primary secondary outcomes.'),
 ('27. Why might retry fail?','The same agent, context, prompt, and corrupted state can reproduce a correlated error. Rollback or independent re-execution changes more of the causal conditions.'),
 ('28. Why a human gate?','Only as a bounded high-risk baseline. It estimates residual value and cost of escalation; it is not treated as a scalable default.'),
 ('29. How do you isolate topology from compute?','Use the same task/seed blocks, model family, information, validator, and total token/tool/timeout caps in the primary comparison. Report a natural-budget analysis separately.'),
 ('30. Why include a single-agent baseline?','It tests whether collaboration adds reliability or merely adds communication surfaces and cost. Multi-agent claims require a strong budget-aware baseline.'),
 ('31. What is the unit of statistical inference?','The task instance. Repeated seeds and multiple faults are nested within tasks and are not treated as independent task samples.'),
 ('32. Why is 100 runs not automatically enough?','Power depends on baseline rate, paired correlation, effect size, and clustering. One hundred paired episode IDs is a floor; pilot-based simulation determines the confirmatory total.'),
 ('33. Which statistical tests will you use?','Exact McNemar for paired binary outcomes, paired permutation or Wilcoxon analyses for skewed paired metrics, task-clustered bootstrap CIs, and mixed models for nested effects.'),
 ('34. How do you control multiple comparisons?','Predeclare primary contrasts, organize them by RQ, and use Holm correction within each confirmatory family.'),
 ('35. Why report effect sizes if p-values are available?','Statistical significance does not show operational importance. Risk difference, matched odds ratio, median shift, or Cliff’s delta quantify magnitude.'),
 ('36. How are timeouts handled?','As explicit outcome categories and normally as task failures, not dropped observations. Provider incidents and exclusions remain in the manifest.'),
 ('37. How does AI Command Center fit?','It launches and observes episodes. MAS-RELIAB owns the scientific task catalog, fault operators, graph/metric definitions, analysis, and report.'),
 ('38. Can MAS-RELIAB work without AI Command Center?','Yes, if another substrate implements the required execution, trace, snapshot, and artifact interfaces. The separation is deliberate.'),
 ('39. What is the biggest threat to validity?','The gap between controlled injected faults and natural deployment failures, combined with incomplete causal visibility. The design addresses but cannot eliminate it.'),
 ('40. How will you handle provider drift?','Pin versions where possible, interleave conditions, log dates/regions/incidents, use snapshots and daily controls, and include time effects in sensitivity analyses.'),
 ('41. What if all hypotheses are rejected?','Report that result. The protocol may still reveal measurement properties, but I would not claim the anticipated effects. Hypotheses are falsifiable, not commitments.'),
 ('42. Why are the results blank?','Because no run data were supplied. Filling them would be fabrication. The locked shell demonstrates exactly what will be reported after execution.'),
 ('43. Is this a completed empirical paper?','No. It is a complete pre-experimental protocol and viva package. Empirical completion requires the frozen runs, validated analysis, and rewritten abstract, Results, discussion, and conclusion.'),
 ('44. What would count as success for the research?','A reproducible execution showing that the measures are computable and reliable, the planned contrasts are adequately powered, and the framework yields actionable distinctions without hiding uncertainty or cost.'),
 ('45. What is the next step?','Freeze task instances and operator doses in a pilot, preregister the manifest and primary analyses, execute interleaved paired runs, validate denominators, and populate the locked tables automatically.')
]
for q,a in qa:
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(2); r=p.add_run(q); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
    p=d.add_paragraph(a); p.paragraph_format.left_indent=Inches(.22); p.paragraph_format.space_after=Pt(5)

d.add_heading('V5. Defence boundaries — phrases to use and avoid',level=2)
table(d,['Use','Avoid'],[
 ('“MAS-RELIAB specifies a unified experimental methodology.”','“This is the first framework for multi-agent reliability.”'),
 ('“The protocol will test whether…”','“The results prove…” before executions exist.'),
 ('“From Spark to Fire already studies graph-based cascades and mitigation.”','Implying graph propagation is novel.'),
 ('“Metrics are reported separately; Pareto analysis shows trade-offs.”','Calling an arbitrary weighted average scientifically objective.'),
 ('“Valid, consumed injections define the propagation denominator.”','Counting every scheduled injection as successful.'),
 ('“State-based validators are primary; model judges are secondary.”','Relying entirely on subjective text matching.'),
 ('“Unsupported hypotheses will be rejected or left inconclusive.”','Reframing contradictions as confirmations.')],widths=[3.25,3.25],small=8.2)

d.add_heading('V6. Final 30-second closing answer',level=2)
para(d,'MAS-RELIAB is valuable because it turns a vague statement—“this multi-agent system is reliable”—into a reproducible set of questions with explicit denominators, interventions, graphs, evidence views, costs, and statistical decisions. It does not replace prior benchmarks or claim priority for cascades. Its contribution is to connect reliability, propagation, attribution, mitigation, topology, and cost in the same auditable protocol. The next scientific obligation is execution, not stronger wording.')

# Save DOCX
out=ROOT/'MAS_RELIAB_Research_Paper_and_Viva_Dev_Parth.docx'
d.save(out)

# Companion schemas and templates
schema_task={
 '$schema':'https://json-schema.org/draft/2020-12/schema','title':'MAS-RELIAB Task','type':'object',
 'required':['task_id','task_version','family','instruction','initial_state_digest','gold','split'],
 'properties':{
  'task_id':{'type':'string'},'task_version':{'type':'string'},'family':{'type':'string'},'instruction':{'type':'string'},
  'initial_state_digest':{'type':'string'},'tools':{'type':'array'},'gold':{'type':'object'},'budgets':{'type':'object'},
  'injection_points':{'type':'array'},'split':{'enum':['dev','pilot','locked']},'provenance':{'type':'object'},'license':{'type':'string'}}}
(SUP/'MAS_RELIAB_task_schema.json').write_text(json.dumps(schema_task,indent=2),encoding='utf-8')

manifest_headers=['episode_id','block_id','task_id','task_version','task_family','split','architecture','topology','model_provider','model_version','prompt_hash','environment_digest','tool_versions','seed','fault_operator','fault_target','fault_position','fault_severity','fault_seed','verification_regime','recovery_strategy','observability_view','token_budget','tool_call_budget','timeout_seconds','schedule_order','manifest_hash']
with open(SUP/'MAS_RELIAB_experiment_manifest_template.csv','w',newline='',encoding='utf-8') as f: csv.writer(f).writerow(manifest_headers)
result_headers=['episode_id','run_id','status','eligible','exclusion_reason','injection_applied','injection_consumed','task_success','constraint_success_json','epr_episode','reachable_downstream','affected_downstream','daf','amplification','propagation_depth','detected','time_to_detect_ms','attributed_agent','gold_agent','attributed_step','gold_step','attribution_correct_agent','attribution_correct_step','recovered','input_tokens','output_tokens','tool_calls','latency_ms','cost_currency','cost_value','trace_digest','validator_version','analysis_version']
with open(SUP/'MAS_RELIAB_results_template.csv','w',newline='',encoding='utf-8') as f: csv.writer(f).writerow(result_headers)

print(out)
